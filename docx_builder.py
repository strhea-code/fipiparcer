"""Генератор docx-карточек для печати.

Альбомная A4, 2 колонки. Каждая карточка — отдельная редактируемая
таблица Word 1×1. Внутренние линии сделаны границами абзацев, поэтому
карточка не распадается между колонками/страницами и её удобно двигать целиком.
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree

from extractor import Task, Chunk, clean_chunks
from math_convert import fipi_mathml_to_omath

OMATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

A4_LANDSCAPE_WIDTH_CM = 29.7
A4_LANDSCAPE_HEIGHT_CM = 21.0
PAGE_MARGIN_CM = 1.0
CARD_WIDTH_CM = 13.0
SECTION_GAP_CM = 0.5

# 0 = автоматическая высота по контенту; 4/6 = ориентир для печати на A4.
CARD_HEIGHT_BY_PER_PAGE: dict[int, float | None] = {0: None, 4: 9.0, 6: 6.0}
# Крупные задания с несколькими картинками не должны разрывать карточку между страницами.
# Поэтому для задач с 2+ изображениями дополнительно ограничиваем высоту каждой картинки.
MULTI_IMAGE_MAX_HEIGHT_CM = 4.1
SINGLE_IMAGE_MAX_HEIGHT_CM = 6.2
ASSETS_DIR = Path(__file__).parent / "assets"


def _setup_landscape_a4(section, margin_cm: float = PAGE_MARGIN_CM) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(A4_LANDSCAPE_WIDTH_CM)
    section.page_height = Cm(A4_LANDSCAPE_HEIGHT_CM)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)


def _set_section_columns(section, num: int = 2, gap_cm: float = SECTION_GAP_CM) -> None:
    """Две колонки: отдельные таблицы-карточки текут сверху вниз и затем вправо."""
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:cols")):
        sectPr.remove(old)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(int(gap_cm * 567)))
    sectPr.append(cols)


def _set_table_fixed_width(table, width_cm: float) -> None:
    width_dxa = int(width_cm * 567)
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(width_dxa))
    tblPr.append(tblW)
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for col in table.columns:
        col.width = Cm(width_cm)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(width_cm)


def _set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_cell_margins(cell, top=0.07, start=0.12, bottom=0.07, end=0.12) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(int(v * 567)))
        node.set(qn("w:type"), "dxa")


def _keep_table_together(table) -> None:
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        trPr = row._tr.get_or_add_trPr()
        for old in trPr.findall(qn("w:cantSplit")):
            trPr.remove(old)
        trPr.append(OxmlElement("w:cantSplit"))


def _set_min_row_height(table, height_cm: float) -> None:
    for row in table.rows:
        row.height = Cm(height_cm)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _paragraph_border(paragraph, where: str = "bottom") -> None:
    """Линия внутри карточки: визуально как граница строки, но без риска разрыва таблицы."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    border = pBdr.find(qn(f"w:{where}"))
    if border is None:
        border = OxmlElement(f"w:{where}")
        pBdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "6")
    border.set(qn("w:space"), "2")
    border.set(qn("w:color"), "000000")


def _format_paragraph(paragraph, font_size_pt: float = 10.5) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size_pt)


def _clear_cell(cell) -> None:
    for p in list(cell.paragraphs):
        p._p.getparent().remove(p._p)


def _add_text_or_math_to_paragraph(paragraph, ch: Chunk) -> None:
    if ch.kind == "text":
        if ch.value:
            run = paragraph.add_run(ch.value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10.5)
    elif ch.kind == "math":
        if not ch.mathml:
            return
        try:
            # В норме extractor отдаёт настоящий MathML из mjx-assistive-mml.
            # Если MathML недоступен и пришёл TeX-текст, не теряем формулу,
            # а вставляем её как редактируемый текст.
            if not ch.mathml.lstrip().startswith("<"):
                run = paragraph.add_run(ch.mathml)
                run.font.name = "Times New Roman"
                run.font.size = Pt(10.5)
                return
            omath_xml = fipi_mathml_to_omath(ch.mathml)
            wrapped = f'<root xmlns:m="{OMATH_NS}">{omath_xml}</root>'
            root = etree.fromstring(wrapped.encode("utf-8"))
            for elem in root:
                paragraph._p.append(elem)
        except Exception as e:
            paragraph.add_run(ch.mathml if ch.mathml else " [formula?] ")
            print(f"[builder] math fail: {e}", file=sys.stderr)
    elif ch.kind == "break":
        paragraph.add_run().add_break()


def _add_chunks_with_images(cell, chunks: list[Chunk], max_img_cm: float, max_img_height_cm: float | None = None) -> None:
    current = cell.add_paragraph()
    _format_paragraph(current)
    for ch in chunks:
        if ch.kind == "image":
            if not ch.image_bytes:
                continue
            p_img = cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(2)
            p_img.paragraph_format.space_after = Pt(2)
            run = p_img.add_run()
            try:
                img_w_cm, img_h_cm = _fit_image_size_cm(ch.image_bytes, max_img_cm, max_img_height_cm)
                if img_h_cm:
                    run.add_picture(io.BytesIO(ch.image_bytes), width=Cm(img_w_cm), height=Cm(img_h_cm))
                else:
                    run.add_picture(io.BytesIO(ch.image_bytes), width=Cm(img_w_cm))
            except Exception as e:
                print(f"[builder] image fail: {e}", file=sys.stderr)
            current = cell.add_paragraph()
            _format_paragraph(current)
        else:
            _add_text_or_math_to_paragraph(current, ch)


def _fit_image_size_cm(image_bytes: bytes, max_width_cm: float, max_height_cm: float | None = None) -> tuple[float, float | None]:
    """Вернуть размеры картинки в сантиметрах с сохранением пропорций.

    Ограничение по высоте важно для задач вроде FADB4B, где несколько графиков
    иначе растягивают карточку на весь лист и Word переносит нижнюю часть на
    следующую страницу.
    """
    try:
        from PIL import Image as PILImage
        img = PILImage.open(io.BytesIO(image_bytes))
        native_w_cm = img.width / 96 * 2.54
        native_h_cm = img.height / 96 * 2.54
        scale = min(1.0, max_width_cm / native_w_cm)
        if max_height_cm:
            scale = min(scale, max_height_cm / native_h_cm)
        return native_w_cm * scale, native_h_cm * scale
    except Exception:
        return max_width_cm, None


def _count_images(chunks: list[Chunk]) -> int:
    return sum(1 for ch in chunks if ch.kind == "image" and ch.image_bytes)


def is_large_task(task: Task, min_images: int = 2, min_text_chars: int = 900) -> bool:
    """Вернуть True для карточек, которые лучше вынести в отдельный файл.

    Типичный пример — задания с несколькими графиками/рисунками. Если класть
    их в общий файл на 6 карточек на лист, Word разрывает карточку между
    страницами или растягивает её на весь лист.
    """
    cleaned = clean_chunks(task.content)
    image_count = _count_images(cleaned)
    text_len = len(" ".join(ch.value for ch in cleaned if ch.kind == "text"))
    return image_count >= min_images or (image_count >= 1 and text_len >= min_text_chars)


def _fill_card(cell, task: Task, with_answer_squares: bool, max_img_cm: float) -> None:
    _clear_cell(cell)
    _set_cell_margins(cell)

    # Заголовок с номером сверху и линией под ним.
    p_num = cell.add_paragraph()
    p_num.paragraph_format.space_before = Pt(0)
    p_num.paragraph_format.space_after = Pt(2)
    run = p_num.add_run(f"Номер: {task.qid}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    _paragraph_border(p_num, "bottom")

    # Тело задания: редактируемый текст, OMath-формулы, картинки.
    cleaned = clean_chunks(task.content)
    image_count = _count_images(cleaned)
    max_img_height_cm = None
    if image_count >= 2:
        max_img_height_cm = MULTI_IMAGE_MAX_HEIGHT_CM
    elif image_count == 1:
        max_img_height_cm = SINGLE_IMAGE_MAX_HEIGHT_CM
    _add_chunks_with_images(cell, cleaned, max_img_cm, max_img_height_cm=max_img_height_cm)

    # Ответ с квадратиками в одной строке, как в ручном образце.
    if with_answer_squares:
        p_ans = cell.add_paragraph()
        p_ans.paragraph_format.space_before = Pt(4)
        p_ans.paragraph_format.space_after = Pt(0)
        _paragraph_border(p_ans, "top")
        run_lbl = p_ans.add_run("Ответ:   ")
        run_lbl.bold = True
        run_lbl.font.name = "Times New Roman"
        run_lbl.font.size = Pt(11)
        run_img = p_ans.add_run()
        try:
            run_img.add_picture(str(ASSETS_DIR / "answer_squares.png"), width=Cm(7.6))
        except Exception as e:
            print(f"[builder] answer-squares fail: {e}", file=sys.stderr)


def build_docx(
    tasks: list[Task],
    out_path: Path,
    with_answer_squares: bool = True,
    card_width_cm: float = CARD_WIDTH_CM,
    per_page: int = 0,
) -> None:
    """Собрать docx: каждая карточка = отдельная таблица 1×1.

    per_page: 0 = авто, 4 = минимум 9 см, 6 = минимум 6 см.
    Развёрнутые задачи можно собирать без поля ответа: with_answer_squares=False.
    """
    doc = Document()
    section = doc.sections[0]
    _setup_landscape_a4(section)
    _set_section_columns(section, num=2)

    if not tasks:
        doc.add_paragraph("Пусто.")
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return

    card_height_cm = CARD_HEIGHT_BY_PER_PAGE.get(per_page, None)

    for task in tasks:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        table.autofit = False
        _set_table_fixed_width(table, card_width_cm)
        _set_table_borders(table)
        _keep_table_together(table)
        if card_height_cm:
            _set_min_row_height(table, card_height_cm)
        _fill_card(table.cell(0, 0), task, with_answer_squares, max_img_cm=card_width_cm - 0.6)
        sep = doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(0)
        sep.paragraph_format.space_after = Pt(1)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(
        f"[builder] {out_path} — {len(tasks)} карточек, "
        f"ширина {card_width_cm} см, высота {'авто' if not card_height_cm else f'{card_height_cm} см'}",
        file=sys.stderr,
    )


def group_tasks(tasks: list[Task]) -> list[Task]:
    """Группировка: похожие начальные формулировки рядом, развёрнутые в конец."""
    short_and_choice = [t for t in tasks if t.answer_type != "extended"]
    extended = [t for t in tasks if t.answer_type == "extended"]
    short_and_choice.sort(key=lambda t: t.first_words)
    extended.sort(key=lambda t: t.first_words)
    return short_and_choice + extended


def split_by_answer_type(tasks: list[Task]) -> tuple[list[Task], list[Task]]:
    short = [t for t in tasks if t.answer_type != "extended"]
    extended = [t for t in tasks if t.answer_type == "extended"]
    return short, extended


def split_large_tasks(tasks: list[Task]) -> tuple[list[Task], list[Task]]:
    """Разделить короткие карточки на обычные и крупные.

    Крупные карточки лучше сохранять в отдельный docx, чтобы они не ломали
    раскладку основного файла с 4–6 карточками на лист.
    """
    regular: list[Task] = []
    large: list[Task] = []
    for task in tasks:
        if is_large_task(task):
            large.append(task)
        else:
            regular.append(task)
    return regular, large
